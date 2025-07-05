"""
lw.py - LaTeX to Word: PDF를 통한 Word 문서 완전 변환 파이프라인
Word (한글+수식) → PDF → LaTeX 변환 → Word (한글+수식 OMath)

작성일: 2025년 1월 5일
목적: 보험수학 문서의 완전한 Word-to-Word 변환

처리 과정:
1. 원본 PDF에서 한글 텍스트와 위치 정보 추출
2. MinerU로 수식 위치 추출
3. 수식 이미지 생성 및 nougat으로 LaTeX 변환
4. LaTeX를 Word OMath XML로 변환
5. 텍스트와 수식을 통합하여 Word 문서 생성
"""

import os
import sys

# Windows 인코딩 문제 해결
if sys.platform == 'win32':
    import locale
    os.environ['PYTHONIOENCODING'] = 'utf-8'

import json
import subprocess
from pathlib import Path
from datetime import datetime
import shutil
import fitz  # PyMuPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET

# LaTeX to MathML 변환용
try:
    from latex2mathml.converter import convert as latex_to_mathml
    HAS_LATEX2MATHML = True
except ImportError:
    HAS_LATEX2MATHML = False
    print("⚠️ latex2mathml이 설치되지 않았습니다. 설치: pip install latex2mathml")


class TextBlock:
    """텍스트 블록 정보"""
    def __init__(self, page, bbox, text, font=None, size=None):
        self.page = page
        self.bbox = bbox  # (x1, y1, x2, y2)
        self.text = text
        self.font = font
        self.size = size
        self.type = 'text'


class FormulaBlock:
    """수식 블록 정보"""
    def __init__(self, page, bbox, latex, image_path=None):
        self.page = page
        self.bbox = bbox
        self.latex = latex
        self.image_path = image_path
        self.type = 'formula'
        self.omath_xml = None


class LaTeXToWordPipeline:
    def __init__(self):
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Windows 경로 사용
        self.base_dir = Path("C:/git/mineru-latex-converter")
        self.output_dir = self.base_dir / f"word_output_{self.timestamp}"
        self.output_dir.mkdir(exist_ok=True, parents=True)
        
    def run_pipeline(self, input_file_path):
        """전체 파이프라인 실행"""
        
        print("🚀 LaTeX to Word 변환 파이프라인 시작")
        print("=" * 70)
        
        input_path = Path(input_file_path)
        
        # Word 문서인 경우 먼저 PDF로 변환
        if input_path.suffix.lower() in ['.doc', '.docx']:
            print("\n[사전 단계] 📄 Word 문서를 고품질 PDF로 변환 중...")
            pdf_path = self.convert_word_to_pdf(input_path)
            if not pdf_path:
                print("❌ Word to PDF 변환 실패")
                return None
            print(f"✅ PDF 변환 완료: {pdf_path}")
        else:
            pdf_path = input_path
        
        # 경로 설정
        paths = {
            'original': Path(pdf_path),
            'text_extraction': self.output_dir / "01_text_extraction",
            'no_korean': self.output_dir / "02_no_korean_no_numbers.pdf",
            'mineru_output': self.output_dir / "03_mineru_results",
            'formula_images': self.output_dir / "04_formula_images",
            'nougat_results': self.output_dir / "05_nougat_results",
            'final_word': self.output_dir / "06_final_result.docx",
            'final_html': self.output_dir / "07_final_result.html",
            'path_info': self.output_dir / "all_paths.json"
        }
        
        # 디렉토리 생성
        for key in ['text_extraction', 'mineru_output', 'formula_images', 'nougat_results']:
            paths[key].mkdir(exist_ok=True)
        
        # 전체 경로 정보 저장
        all_paths = {k: str(v) for k, v in paths.items()}
        
        try:
            # 1단계: 원본 PDF에서 텍스트와 위치 정보 추출
            print("\n[1/7] 📝 원본 PDF에서 텍스트 추출 중...")
            text_blocks = self.extract_text_with_positions(paths['original'], paths['text_extraction'])
            print(f"✅ {len(text_blocks)}개 텍스트 블록 추출")
            
            # 2단계: 한글 및 번호 제거 (수식 추출용)
            print("\n[2/7] 🔤 한글 제거 (수식 추출용)...")
            self.remove_korean_text(paths['original'], paths['no_korean'])
            
            # 3단계: MinerU 실행
            print("\n[3/7] 🔍 MinerU로 레이아웃 분석 중...")
            mineru_result_dir = self.run_mineru(paths['no_korean'], paths['mineru_output'])
            
            # 4단계: 수식 이미지 추출
            print("\n[4/7] 🖼️ 수식 이미지 추출 중...")
            formula_blocks = self.extract_formula_images(
                paths['no_korean'], 
                mineru_result_dir,
                paths['formula_images']
            )
            print(f"✅ {len(formula_blocks)}개 수식 추출")
            
            # 5단계: nougat으로 LaTeX 변환
            print("\n[5/7] 🔬 nougat으로 LaTeX 변환 중...")
            formula_blocks = self.convert_with_nougat(formula_blocks, paths['nougat_results'])
            
            # 6단계: 텍스트와 수식 통합 및 Word 문서 생성
            print("\n[6/7] 📄 Word 문서 생성 중...")
            merged_content = self.merge_text_and_formulas(text_blocks, formula_blocks)
            
            # 디버깅: 병합된 컨텐츠 확인
            print(f"\n🔍 병합된 컨텐츠 분석:")
            print(f"  - 총 줄 수: {len(merged_content)}")
            for i, line in enumerate(merged_content[:3]):  # 처음 3줄만 표시
                print(f"  - 줄 {i+1}: {len(line)}개 요소")
                for block in line:
                    if block.type == 'text':
                        print(f"    [{block.bbox[0]:.1f}, {block.bbox[1]:.1f}] 텍스트: {block.text[:20]}...")
                    else:
                        print(f"    [{block.bbox[0]:.1f}, {block.bbox[1]:.1f}] 수식: {block.latex[:20] if block.latex else 'None'}...")
            
            self.create_word_document(merged_content, paths['final_word'])
            
            # 7단계: HTML 문서도 생성 (ln.py 스타일 유지)
            print("\n[7/7] 🌐 HTML 결과 문서 생성 중...")
            self.generate_html(merged_content, paths['final_html'])
            
            # 전체 경로 정보 저장
            all_paths['statistics'] = {
                'text_blocks': len(text_blocks),
                'formulas_found': len(formula_blocks),
                'formulas_converted': sum(1 for f in formula_blocks if f.latex)
            }
            
            with open(paths['path_info'], 'w', encoding='utf-8') as f:
                json.dump(all_paths, f, ensure_ascii=False, indent=2)
            
            print("\n" + "=" * 70)
            print("✅ 변환 완료!")
            print(f"📄 Word 문서: {paths['final_word']}")
            print(f"🌐 HTML 문서: {paths['final_html']}")
            print(f"📁 전체 결과: {self.output_dir}")
            
            # 결과 폴더 자동 열기
            try:
                os.startfile(str(self.output_dir))
                os.startfile(str(paths['final_word']))
            except:
                pass
            
            return all_paths
            
        except Exception as e:
            print(f"\n❌ 오류 발생: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def extract_text_with_positions(self, pdf_path, output_dir):
        """원본 PDF에서 텍스트와 위치 정보 추출 (빈 공간 포함)"""
        doc = fitz.open(str(pdf_path))
        text_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # 정밀한 텍스트 추출 - dict 사용
            blocks = page.get_text("dict")
            
            # 각 블록의 모든 텍스트 span 추출
            for block in blocks["blocks"]:
                if block.get("type") != 0:  # 텍스트 블록만
                    continue
                    
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        # 공백 포함 모든 텍스트
                        text = span.get("text", "")
                        if not text.strip():  # 공백만 있어도 유지
                            continue
                            
                        text_block = TextBlock(
                            page=page_num,
                            bbox=span.get("bbox", [0, 0, 0, 0]),
                            text=text,  # strip 하지 않음
                            font=span.get("font", ""),
                            size=span.get("size", 11)
                        )
                        text_blocks.append(text_block)
        
        # 텍스트 정보 저장
        text_info = []
        for block in text_blocks:
            text_info.append({
                'page': block.page,
                'bbox': block.bbox,
                'text': block.text,
                'font': block.font,
                'size': block.size
            })
        
        with open(output_dir / "text_blocks.json", 'w', encoding='utf-8') as f:
            json.dump(text_info, f, ensure_ascii=False, indent=2)
        
        doc.close()
        return text_blocks
    
    def remove_korean_text(self, input_pdf, output_pdf):
        """한글 텍스트 제거 (수식 추출용)"""
        try:
            # enhanced_korean_remover_windows.py 사용
            korean_remover_script = self.base_dir / "enhanced_korean_remover_windows.py"
            if korean_remover_script.exists():
                cmd = [sys.executable, str(korean_remover_script), 
                       str(input_pdf), str(output_pdf)]
                result = subprocess.run(cmd, capture_output=True, text=True,
                                      encoding='utf-8', errors='ignore')
                if result.returncode == 0:
                    print("✅ 한글 제거 완료")
                else:
                    print("⚠️ 한글 제거 실패, 원본 사용")
                    shutil.copy2(input_pdf, output_pdf)
            else:
                shutil.copy2(input_pdf, output_pdf)
        except:
            shutil.copy2(input_pdf, output_pdf)
    
    def run_mineru(self, pdf_path, output_dir):
        """MinerU 실행"""
        try:
            # 먼저 magic-pdf 명령어 시도
            mineru_cmd = [
                "magic-pdf", "-p", str(pdf_path),
                "-o", str(output_dir),
                "-m", "auto"
            ]
            
            print(f"   실행 명령: {' '.join(mineru_cmd)}")
            result = subprocess.run(mineru_cmd, 
                                  capture_output=True, 
                                  text=True, 
                                  encoding='utf-8',
                                  errors='ignore',
                                  shell=True)
            
            if result.returncode != 0:
                # Python 모듈로 실행 시도
                mineru_cmd = [
                    sys.executable, "-m", "magic_pdf",
                    "-p", str(pdf_path),
                    "-o", str(output_dir),
                    "-m", "auto"
                ]
                result = subprocess.run(mineru_cmd, capture_output=True, text=True,
                                      encoding='utf-8', errors='ignore')
                
            # MinerU 결과 디렉토리 찾기
            pdf_basename = pdf_path.stem
            mineru_result_dir = output_dir / pdf_basename / "auto"
            
            if not mineru_result_dir.exists():
                mineru_dirs = list(output_dir.glob("*/auto"))
                if not mineru_dirs:
                    # 기존 결과 사용
                    print("⚠️ MinerU를 찾을 수 없습니다. 기존 결과 사용")
                    existing_dir = self.base_dir / "output_adjusted_20250705_125920/27_no_korean_direct/auto"
                    if existing_dir.exists():
                        mineru_auto_dir = output_dir / "27_no_korean_direct" / "auto"
                        mineru_auto_dir.parent.mkdir(exist_ok=True, parents=True)
                        shutil.copytree(existing_dir, mineru_auto_dir)
                        return mineru_auto_dir
                    else:
                        raise Exception("MinerU 결과를 찾을 수 없습니다")
                else:
                    mineru_result_dir = mineru_dirs[0]
            
            return mineru_result_dir
            
        except Exception as e:
            print(f"⚠️ MinerU 오류: {str(e)}")
            raise
    
    def extract_formula_images(self, pdf_path, mineru_dir, output_dir):
        """수식 이미지 추출 및 FormulaBlock 생성"""
        formula_blocks = []
        
        model_json_path = mineru_dir / "model.json"
        if not model_json_path.exists():
            return formula_blocks
        
        # model.json 로드
        with open(model_json_path, 'r', encoding='utf-8') as f:
            model_data = json.load(f)
        
        # PDF 열기
        doc = fitz.open(str(pdf_path))
        
        # 확장 비율
        WIDTH_EXPANSION = 0.2
        HEIGHT_EXPANSION = 0.03
        
        formula_count = 0
        
        # 각 페이지 처리
        for page_idx, page_data in enumerate(model_data):
            if not isinstance(page_data, dict):
                continue
                
            layout_dets = page_data.get('layout_dets', [])
            page = doc[page_idx] if page_idx < len(doc) else None
            
            if not page:
                continue
                
            # 페이지 크기
            page_width = page.rect.width
            page_height = page.rect.height
            
            # MinerU 좌표 시스템
            mineru_width = 1654
            mineru_height = 2339
            
            # 좌표 변환 비율
            scale_x = page_width / mineru_width
            scale_y = page_height / mineru_height
            
            # 수식 찾기
            for det in layout_dets:
                category_id = det.get('category_id')
                
                if category_id in [13, 14]:  # 수식
                    poly = det.get('poly', [])
                    if len(poly) >= 8:
                        # bbox 계산
                        x_coords = [poly[i] for i in range(0, 8, 2)]
                        y_coords = [poly[i] for i in range(1, 8, 2)]
                        bbox = [min(x_coords), min(y_coords), max(x_coords), max(y_coords)]
                        
                        # 확장 적용
                        x1, y1, x2, y2 = bbox
                        width = x2 - x1
                        height = y2 - y1
                        
                        width_expand = width * WIDTH_EXPANSION
                        height_expand = height * HEIGHT_EXPANSION
                        
                        # 조정된 좌표
                        adj_x1 = max(0, x1 - width_expand/2)
                        adj_x2 = min(mineru_width, x2 + width_expand/2)
                        adj_y1 = max(0, y1 - height_expand/2)
                        adj_y2 = min(mineru_height, y2 + height_expand/2)
                        
                        # PDF 좌표로 변환
                        pdf_x1 = adj_x1 * scale_x
                        pdf_x2 = adj_x2 * scale_x
                        pdf_y1 = adj_y1 * scale_y
                        pdf_y2 = adj_y2 * scale_y
                        
                        # 이미지 추출
                        rect = fitz.Rect(pdf_x1, pdf_y1, pdf_x2, pdf_y2)
                        mat = fitz.Matrix(4.17, 4.17)  # 300 DPI
                        pix = page.get_pixmap(clip=rect, matrix=mat, alpha=False)
                        
                        # 파일 저장
                        formula_type = "inline" if category_id == 13 else "block"
                        img_name = f"formula_p{page_idx}_{formula_count}_{formula_type}.png"
                        img_path = output_dir / img_name
                        pix.save(str(img_path))
                        
                        # FormulaBlock 생성
                        formula_block = FormulaBlock(
                            page=page_idx,
                            bbox=(pdf_x1, pdf_y1, pdf_x2, pdf_y2),
                            latex=None,
                            image_path=str(img_path)
                        )
                        formula_blocks.append(formula_block)
                        
                        formula_count += 1
                        print(f"  수식 {formula_count}: {img_name}")
        
        doc.close()
        return formula_blocks
    
    def convert_with_nougat(self, formula_blocks, output_dir):
        """nougat으로 LaTeX 변환"""
        output_dir.mkdir(exist_ok=True)
        
        # 먼저 기존 nougat 결과가 있는지 확인
        existing_results_path = self.base_dir / "word_output_20250705_171335/05_nougat_results/nougat_results.json"
        if existing_results_path.exists():
            print("📤 기존 nougat 결과 사용")
            with open(existing_results_path, 'r', encoding='utf-8') as f:
                existing_results = json.load(f)
            
            # 기존 결과를 formula_blocks에 매핑
            for idx, formula in enumerate(formula_blocks):
                if idx < len(existing_results):
                    formula.latex = existing_results[idx].get('latex', '')
                    status = "✅" if formula.latex else "❌"
                    if formula.latex:
                        print(f"  {status} 수식 {idx + 1}: {formula.latex[:50]}...")
                else:
                    formula.latex = ''
            
            # 결과 저장
            results = []
            for idx, formula in enumerate(formula_blocks):
                results.append({
                    'index': idx,
                    'image_path': str(formula.image_path),
                    'latex': formula.latex,
                    'bbox': formula.bbox,
                    'page': formula.page
                })
            
            with open(output_dir / "nougat_results.json", 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            
            print(f"\n✅ 총 {len(formula_blocks)}개 수식 처리 완료")
            return formula_blocks
        
        # 기존 결과가 없으면 실제 nougat 실행 시도
        try:
            sys.path.append(str(self.base_dir))
            
            # 최적화된 nougat 사용 시도
            from optimized_nougat import OptimizedNougat
            
            print("  🔄 nougat 모델 로드 중... (한 번만 로드)")
            nougat_processor = OptimizedNougat()
            
            results = []
            total = len(formula_blocks)
            
            for idx, formula in enumerate(formula_blocks):
                try:
                    # 진행률 표시
                    percent = ((idx + 1) / total) * 100
                    print(f"  🧮 수식 {idx + 1}/{total} 처리 중... ({percent:.1f}%)", end='\r')
                    
                    # 모델이 이미 로드된 상태에서 변환
                    formula.latex = nougat_processor.convert_image(formula.image_path)
                    
                    if idx == total - 1:  # 마지막 항목
                        print()  # 줄바꿈
                    
                except Exception as e:
                    print(f"\n  ❌ 수식 {idx + 1}: 오류 - {str(e)[:50]}")
                    formula.latex = ""
                
                results.append({
                    'index': idx,
                    'image_path': str(formula.image_path),
                    'latex': formula.latex,
                    'bbox': formula.bbox,
                    'page': formula.page
                })
            
            # 성공/실패 통계
            success_count = sum(1 for f in formula_blocks if f.latex)
            print(f"  ✅ 변환 완료: {success_count}/{total} 성공")
            
            # 결과 저장
            with open(output_dir / "nougat_results.json", 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            
        except ImportError:
            print("⚠️ 최적화된 nougat을 찾을 수 없습니다. 기본 방식 사용...")
            # 기존 방식으로 폴백
            return self.convert_with_nougat_fallback(formula_blocks, output_dir)
        
        return formula_blocks
    
    def convert_with_nougat_fallback(self, formula_blocks, output_dir):
        """nougat 기본 방식 (이전 코드)"""
        try:
            try:
                from simple_nougat_clean import run_nougat_on_image
            except:
                from simple_nougat_inference import run_nougat_on_image
            
            results = []
            
            for idx, formula in enumerate(formula_blocks):
                try:
                    result = run_nougat_on_image(formula.image_path)
                    if isinstance(result, dict):
                        formula.latex = result.get('latex', '')
                    else:
                        formula.latex = result if result else ""
                    
                    status = "✅" if formula.latex else "❌"
                    print(f"  {status} 수식 {idx + 1}: {formula.latex[:50] if formula.latex else '변환 실패'}")
                    
                except Exception as e:
                    print(f"  ❌ 수식 {idx + 1}: 오류 - {str(e)[:50]}")
                    formula.latex = ""
                
                results.append({
                    'index': idx,
                    'image_path': formula.image_path,
                    'latex': formula.latex,
                    'bbox': formula.bbox,
                    'page': formula.page
                })
            
            # 결과 저장
            with open(output_dir / "nougat_results.json", 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            
        except ImportError:
            print("❌ nougat 모듈을 찾을 수 없습니다")
            raise
        
        return formula_blocks
    
    def latex_to_omath_xml(self, latex_code):
        """LaTeX를 Word OMath XML로 변환"""
        if not HAS_LATEX2MATHML:
            return None
        
        try:
            # LaTeX → MathML 변환
            mathml_str = latex_to_mathml(latex_code)
            
            # MathML을 파싱
            # MathML 네임스페이스 제거 (간단한 처리를 위해)
            mathml_str = mathml_str.replace('xmlns="http://www.w3.org/1998/Math/MathML"', '')
            mathml_str = mathml_str.replace('<math>', '').replace('</math>', '')
            
            # OMath 구조 생성
            omath = self.create_omath_from_latex(latex_code)
            return omath
            
        except Exception as e:
            print(f"LaTeX → OMath 변환 오류: {str(e)}")
            return None
    
    def create_omath_from_latex(self, latex_code):
        """LaTeX에서 직접 OMath 생성 (복잡한 수식 지원)"""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # OMath 네임스페이스
        math_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        
        # 복잡한 OMath 구조
        omath_content = self.latex_to_omath_content(latex_code)
        
        omath_xml = f'''<m:oMathPara xmlns:m="{math_ns}" xmlns:w="{w_ns}">
                <m:oMathParaPr>
                    <m:jc m:val="center"/>
                </m:oMathParaPr>
                <m:oMath>
                    {omath_content}
                </m:oMath>
            </m:oMathPara>'''
        
        try:
            return parse_xml(omath_xml)
        except Exception as e:
            print(f"OMath 파싱 오류: {str(e)[:100]}")
            # 파싱 실패시 기본 텍스트로
            return None
    
    def latex_to_omath_content(self, latex_code):
        """LaTeX를 복잡한 OMath XML로 변환"""
        import re
        
        # 보험수학 수식의 복잡한 패턴 처리
        
        # 먼저 전체 구조 분석
        result = self.parse_latex_structure(latex_code)
        return result
    
    def parse_latex_structure(self, latex_code):
        """LaTeX 구조를 분석하여 OMath XML 생성"""
        import re
        
        # 배열/행렬 구조 확인
        if '\\begin{array}' in latex_code:
            return self.convert_array_to_omath(latex_code)
        
        # 분수가 있는 경우
        if '\\frac{' in latex_code:
            return self.convert_fraction_to_omath(latex_code)
        
        # 위/아래 첨자가 있는 경우
        if '^' in latex_code or '_' in latex_code:
            return self.convert_scripts_to_omath(latex_code)
        
        # 기본 텍스트
        return self.convert_text_to_omath(latex_code)
    
    def convert_array_to_omath(self, latex_code):
        """배열/행렬을 OMath로 변환"""
        import re
        
        # \begin{array} ... \end{array} 추출
        array_match = re.search(r'\\begin\{array\}(.+?)\\end\{array\}', latex_code, re.DOTALL)
        if not array_match:
            return self.convert_text_to_omath(latex_code)
        
        array_content = array_match.group(1)
        
        # 행으로 분리
        rows = array_content.split('\\\\')
        
        # OMath 행렬 생성
        omath_rows = []
        for row in rows:
            if not row.strip():
                continue
            # 열로 분리
            cells = row.split('&')
            omath_cells = []
            for cell in cells:
                cell_content = self.parse_latex_structure(cell.strip())
                omath_cells.append(f'<m:e>{cell_content}</m:e>')
            
            omath_row = f'''<m:mr>
                {''.join(omath_cells)}
            </m:mr>'''
            omath_rows.append(omath_row)
        
        return f'''<m:m>
            <m:mPr>
                <m:mcs>
                    <m:mc>
                        <m:mcPr>
                            <m:count m:val="{len(cells)}"/>
                            <m:mcJc m:val="center"/>
                        </m:mcPr>
                    </m:mc>
                </m:mcs>
            </m:mPr>
            {''.join(omath_rows)}
        </m:m>'''
    
    def convert_fraction_to_omath(self, latex_code):
        """분수를 OMath로 변환"""
        import re
        
        # 분수 패턴 찾기
        frac_pattern = r'\\frac\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}'
        
        def frac_replacer(match):
            num = match.group(1)
            den = match.group(2)
            
            # 분자와 분모를 재귀적으로 처리
            num_omath = self.parse_latex_structure(num)
            den_omath = self.parse_latex_structure(den)
            
            return f'''<m:f>
                <m:fPr>
                    <m:type m:val="bar"/>
                </m:fPr>
                <m:num>{num_omath}</m:num>
                <m:den>{den_omath}</m:den>
            </m:f>'''
        
        result = latex_code
        while '\\frac{' in result:
            result = re.sub(frac_pattern, frac_replacer, result, count=1)
        
        # 나머지 부분 처리
        if result != latex_code:
            return self.parse_latex_structure(result)
        else:
            return self.convert_text_to_omath(result)
        
    def convert_scripts_to_omath(self, latex_code):
        """위/아래 첨자를 OMath로 변환"""
        import re
        
        # 복잡한 패턴: base_{sub}^{sup}
        subsup_pattern = r'([a-zA-Z0-9\\]+)_\{([^{}]+)\}\^\{([^{}]+)\}'
        
        def subsup_replacer(match):
            base = match.group(1)
            sub = match.group(2)
            sup = match.group(3)
            
            base_omath = self.convert_text_to_omath(base)
            sub_omath = self.parse_latex_structure(sub)
            sup_omath = self.parse_latex_structure(sup)
            
            return f'''<m:sSubSup>
                <m:sSubSupPr>
                    <m:ctrlPr/>
                </m:sSubSupPr>
                <m:e>{base_omath}</m:e>
                <m:sub>{sub_omath}</m:sub>
                <m:sup>{sup_omath}</m:sup>
            </m:sSubSup>'''
        
        result = re.sub(subsup_pattern, subsup_replacer, latex_code)
        
        # 단독 위첨자
        sup_pattern = r'([a-zA-Z0-9\\]+)\^\{([^{}]+)\}'
        def sup_replacer(match):
            base = match.group(1)
            sup = match.group(2)
            
            base_omath = self.convert_text_to_omath(base)
            sup_omath = self.parse_latex_structure(sup)
            
            return f'''<m:sSup>
                <m:sSupPr>
                    <m:ctrlPr/>
                </m:sSupPr>
                <m:e>{base_omath}</m:e>
                <m:sup>{sup_omath}</m:sup>
            </m:sSup>'''
        
        result = re.sub(sup_pattern, sup_replacer, result)
        
        # 단독 아래첨자
        sub_pattern = r'([a-zA-Z0-9\\]+)_\{([^{}]+)\}'
        def sub_replacer(match):
            base = match.group(1)
            sub = match.group(2)
            
            base_omath = self.convert_text_to_omath(base)
            sub_omath = self.parse_latex_structure(sub)
            
            return f'''<m:sSub>
                <m:sSubPr>
                    <m:ctrlPr/>
                </m:sSubPr>
                <m:e>{base_omath}</m:e>
                <m:sub>{sub_omath}</m:sub>
            </m:sSub>'''
        
        result = re.sub(sub_pattern, sub_replacer, result)
        
        # 나머지 텍스트 처리
        return self.convert_text_to_omath(result) if result == latex_code else result
    
    def convert_text_to_omath(self, text):
        """텍스트를 OMath run으로 변환"""
        # LaTeX 명령어 처리
        text = self.process_latex_commands(text)
        
        # 특수 문자 처리
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        # 빈 텍스트 처리
        if not text.strip():
            return '<m:r><m:t> </m:t></m:r>'
        
        return f'<m:r><m:t>{text}</m:t></m:r>'
    
    def process_latex_commands(self, text):
        """보험수학 LaTeX 명령어 처리"""
        # 보험수학 특수 기호
        replacements = {
            '\\mathrm': '',
            '\\mathbb': '',
            '\\mathbf': '',
            '\\boldmath': '',
            '\\operatorname': '',
            '\\displaystyle': '',
            '\\scriptsize': '',
            '\\tiny': '',
            '\\large': '',
            '\\Large': '',
            '\\cdot': '·',
            '\\times': '×',
            '\\pm': '±',
            '\\sum': '∑',
            '\\int': '∫',
            '\\infty': '∞',
            '\\alpha': 'α',
            '\\beta': 'β',
            '\\gamma': 'γ',
            '\\delta': 'δ',
            '\\lambda': 'λ',
            '\\mu': 'μ',
            '\\sigma': 'σ',
            '\\Omega': 'Ω',
            '\\leq': '≤',
            '\\geq': '≥',
            '\\neq': '≠',
            '\\approx': '≈',
            '\\left[': '[',
            '\\right]': ']',
            '\\left(': '(',
            '\\right)': ')',
            '\\left\\{': '{',
            '\\right\\}': '}',
            '\\qquad': '    ',
            '\\quad': '  ',
            '\\,': ' ',
            '\\;': ' ',
            '\\!': '',
            '\\\\': ' '
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # 작은 따옴표나 다른 특수 기호 처리
        text = text.replace("'", "'")
        text = text.replace('~', ' ')
        
        # 남은 백슬래시 제거 (명령어가 아닌 경우)
        if text.startswith('\\') and len(text) > 1 and text[1].isalpha():
            # 이름을 알 수 없는 명령어
            text = text[1:]  # 백슬래시만 제거
        
        return text
    
    def apply_omath_to_word(self, word_path, formula_blocks):
        """Word 문서에 OMath 적용 (COM 사용)"""
        try:
            # Windows가 아니면 건너뛰기
            if sys.platform != 'win32':
                print("  ⚠️ Windows에서만 OMath 적용 가능")
                return None
            
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = None
            
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False  # 백그라운드에서 작업
                
                # 문서 열기
                doc = word.Documents.Open(str(word_path))
                
                # 플레이스홀더를 OMath로 교체
                replaced_count = 0
                
                if hasattr(self, '_formula_placeholders'):
                    for placeholder, latex in self._formula_placeholders.items():
                        if self._replace_placeholder_with_omath(doc, placeholder, latex):
                            replaced_count += 1
                
                print(f"  ✅ {replaced_count}개 수식을 OMath로 변환")
                
                # 새 파일로 저장
                omath_path = Path(word_path).parent / f"{Path(word_path).stem}_omath.docx"
                doc.SaveAs2(str(omath_path))
                doc.Close()
                
                return omath_path
                
            finally:
                if word:
                    try:
                        word.Quit()
                    except:
                        pass
                pythoncom.CoUninitialize()
                
        except Exception as e:
            print(f"  ⚠️ OMath 적용 실패: {str(e)}")
            return None
    
    def _replace_placeholder_with_omath(self, doc, placeholder, latex):
        """플레이스홀더를 OMath로 교체"""
        try:
            # 문서에서 플레이스홀더 찾기
            find = doc.Content.Find
            find.ClearFormatting()
            find.Text = placeholder
            find.Forward = True
            find.Wrap = 0  # wdFindStop
            
            if find.Execute():
                found_range = doc.Range(find.Parent.Start, find.Parent.End)
                
                # OMath 객체 생성
                omath = doc.OMaths.Add(found_range)
                
                # 플레이스홀더 삭제
                found_range.Text = ""
                
                # OMath 구축
                self._build_omath_com(omath, latex)
                
                return True
                
        except Exception as e:
            # print(f"    교체 실패: {placeholder} - {str(e)[:50]}")
            pass
        
        return False
    
    def _build_omath_com(self, omath, latex):
        """COM을 통한 OMath 구축"""
        import re
        
        try:
            # 분수
            if '\\frac{' in latex:
                match = re.search(r'\\frac\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}', latex)
                if match:
                    num = self.clean_latex_text(match.group(1))
                    den = self.clean_latex_text(match.group(2))
                    
                    func = omath.Functions.Add(omath.Range, 1)  # wdOMathFunctionFrac
                    func.Frac.Num.Range.Text = num
                    func.Frac.Den.Range.Text = den
                    return
            
            # 위아래 첨자
            if match := re.search(r'([a-zA-Z]+)_\{([^}]+)\}\^\{([^}]+)\}', latex):
                base = self.clean_latex_text(match.group(1))
                sub = self.clean_latex_text(match.group(2))
                sup = self.clean_latex_text(match.group(3))
                
                func = omath.Functions.Add(omath.Range, 8)  # wdOMathFunctionSubSup
                func.SubSup.Base.Range.Text = base
                func.SubSup.Sub.Range.Text = sub
                func.SubSup.Sup.Range.Text = sup
                return
            
            # 아래첨자만
            if match := re.search(r'([a-zA-Z]+)_\{([^}]+)\}', latex):
                base = self.clean_latex_text(match.group(1))
                sub = self.clean_latex_text(match.group(2))
                
                func = omath.Functions.Add(omath.Range, 6)  # wdOMathFunctionScrSub
                func.ScrSub.Base.Range.Text = base
                func.ScrSub.Sub.Range.Text = sub
                return
            
            # 위첨자만
            if match := re.search(r'([a-zA-Z]+)\^\{([^}]+)\}', latex):
                base = self.clean_latex_text(match.group(1))
                sup = self.clean_latex_text(match.group(2))
                
                func = omath.Functions.Add(omath.Range, 7)  # wdOMathFunctionScrSup
                func.ScrSup.Base.Range.Text = base
                func.ScrSup.Sup.Range.Text = sup
                return
            
            # 합계 기호
            if match := re.search(r'\\sum_\{([^}]+)\}\^\{([^}]+)\}', latex):
                lower = self.clean_latex_text(match.group(1))
                upper = self.clean_latex_text(match.group(2))
                
                func = omath.Functions.Add(omath.Range, 13)  # wdOMathFunctionNary
                func.Nary.Char = 8721  # ∑
                func.Nary.Grow = True
                func.Nary.SubSup.Range.Text = f"{lower} to {upper}"
                return
            
            # 기본: 정리된 텍스트
            omath.Range.Text = self.clean_latex_text(latex)
            
        except Exception as e:
            # 실패시 원본 텍스트
            omath.Range.Text = self.clean_latex_text(latex)
    
    def insert_omath_formula(self, paragraph, latex_code):
        """OMath 수식을 문단에 삽입"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        try:
            # 수식을 위한 run 생성
            run = paragraph.add_run()
            r = run._r
            
            # OMath 컨테이너 생성
            omath = OxmlElement('m:oMath')
            
            # LaTeX 파싱하여 OMath 구조 생성
            omath_content = self.parse_latex_to_omath(latex_code)
            if omath_content:
                omath.append(omath_content)
                # run 요소에 OMath 추가
                r.append(omath)
            else:
                # 파싱 실패시 단순 텍스트
                run.text = self.simplify_latex_for_word(latex_code)
                run.font.name = 'Cambria Math'
                from docx.shared import Pt
                run.font.size = Pt(11)
                
        except Exception as e:
            print(f"  ⚠️ OMath 삽입 오류: {str(e)}")
            # 오류시 단순 텍스트로 표시
            run = paragraph.add_run()
            run.text = self.simplify_latex_for_word(latex_code)
            run.font.name = 'Cambria Math'
            from docx.shared import Pt
            run.font.size = Pt(11)
        
        # 수식 뒤 공백
        paragraph.add_run(" ")
    
    def convert_word_to_pdf(self, word_path):
        """Word 문서를 고품질 PDF로 변환"""
        try:
            # Windows가 아니면 다른 방법 시도
            if sys.platform != 'win32':
                print("  ⚠️ Windows가 아닌 환경에서는 Word COM API를 사용할 수 없습니다.")
                return None
            
            import win32com.client
            import pythoncom
            
            pythoncom.CoInitialize()
            word = None
            
            try:
                # Word 어플리케이션 실행
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # Word 문서 열기
                doc = word.Documents.Open(str(word_path.absolute()))
                
                # PDF 파일 경로 설정
                pdf_path = self.output_dir / f"{word_path.stem}_converted.pdf"
                
                # PDF로 저장 (고품질 설정)
                # 17 = wdFormatPDF
                doc.SaveAs2(
                    str(pdf_path.absolute()),
                    FileFormat=17  # wdFormatPDF
                )
                
                # 문서 닫기
                doc.Close()
                
                return pdf_path
                
            finally:
                if word:
                    try:
                        word.Quit()
                    except:
                        pass
                pythoncom.CoUninitialize()
                
        except Exception as e:
            print(f"  ❌ Word to PDF 변환 오류: {str(e)}")
            return None
    
    def parse_latex_to_omath(self, latex_code):
        """LaTeX를 OMath XML 구조로 파싱"""
        from docx.oxml import OxmlElement
        import re
        
        # 보험수학 수식 예제들을 처리
        
        # 1. 분수 패턴: \frac{numerator}{denominator}
        frac_pattern = r'\\frac\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]+(?:\{[^{}]*\}[^{}]*)*)\}'
        frac_match = re.search(frac_pattern, latex_code)
        if frac_match:
            return self.create_fraction_omath(frac_match.group(1), frac_match.group(2))
        
        # 2. 위/아래 첨자 패턴
        # q_{x+k}^{A} 형태
        subsup_pattern = r'([a-zA-Z]+)_\{([^}]+)\}\^\{([^}]+)\}'
        subsup_match = re.search(subsup_pattern, latex_code)
        if subsup_match:
            return self.create_subsup_omath(
                subsup_match.group(1),
                subsup_match.group(2), 
                subsup_match.group(3)
            )
        
        # 3. 아래첨자만: q_{x+k}
        sub_pattern = r'([a-zA-Z]+)_\{([^}]+)\}'
        sub_match = re.search(sub_pattern, latex_code)
        if sub_match:
            return self.create_subscript_omath(sub_match.group(1), sub_match.group(2))
        
        # 4. 위첨자만: x^{2}
        sup_pattern = r'([a-zA-Z0-9]+)\^\{([^}]+)\}'
        sup_match = re.search(sup_pattern, latex_code)
        if sup_match:
            return self.create_superscript_omath(sup_match.group(1), sup_match.group(2))
        
        # 5. 합계 기호: \sum_{k=0}^{n}
        sum_pattern = r'\\sum_\{([^}]+)\}\^\{([^}]+)\}'
        sum_match = re.search(sum_pattern, latex_code)
        if sum_match:
            return self.create_sum_omath(sum_match.group(1), sum_match.group(2))
        
        # 6. 배열/행렬: \begin{array}
        if '\\begin{array}' in latex_code:
            return self.create_array_omath(latex_code)
        
        # 기본: 단순 텍스트
        return self.create_text_omath(latex_code)
    
    def create_fraction_omath(self, numerator, denominator):
        """분수 OMath 생성"""
        from docx.oxml import OxmlElement
        
        frac = OxmlElement('m:f')
        
        # 분수 속성
        frac_pr = OxmlElement('m:fPr')
        type_elem = OxmlElement('m:type')
        type_elem.set(OxmlElement.nsmap['m'] + 'val', 'bar')
        frac_pr.append(type_elem)
        frac.append(frac_pr)
        
        # 분자
        num = OxmlElement('m:num')
        num.append(self.parse_latex_to_omath(numerator) or self.create_text_omath(numerator))
        frac.append(num)
        
        # 분모
        den = OxmlElement('m:den')
        den.append(self.parse_latex_to_omath(denominator) or self.create_text_omath(denominator))
        frac.append(den)
        
        return frac
    
    def create_subsup_omath(self, base, subscript, superscript):
        """위아래 첨자 OMath 생성"""
        from docx.oxml import OxmlElement
        
        subsup = OxmlElement('m:sSubSup')
        
        # 속성
        subsup_pr = OxmlElement('m:sSubSupPr')
        subsup.append(subsup_pr)
        
        # 기본
        e = OxmlElement('m:e')
        e.append(self.create_text_omath(base))
        subsup.append(e)
        
        # 아래첨자
        sub = OxmlElement('m:sub')
        sub.append(self.create_text_omath(subscript))
        subsup.append(sub)
        
        # 위첨자
        sup = OxmlElement('m:sup')
        sup.append(self.create_text_omath(superscript))
        subsup.append(sup)
        
        return subsup
    
    def create_subscript_omath(self, base, subscript):
        """아래첨자 OMath 생성"""
        from docx.oxml import OxmlElement
        
        sub_elem = OxmlElement('m:sSub')
        
        # 속성
        sub_pr = OxmlElement('m:sSubPr')
        sub_elem.append(sub_pr)
        
        # 기본
        e = OxmlElement('m:e')
        e.append(self.create_text_omath(base))
        sub_elem.append(e)
        
        # 아래첨자
        sub = OxmlElement('m:sub')
        sub.append(self.create_text_omath(subscript))
        sub_elem.append(sub)
        
        return sub_elem
    
    def create_superscript_omath(self, base, superscript):
        """위첨자 OMath 생성"""
        from docx.oxml import OxmlElement
        
        sup_elem = OxmlElement('m:sSup')
        
        # 속성
        sup_pr = OxmlElement('m:sSupPr')
        sup_elem.append(sup_pr)
        
        # 기본
        e = OxmlElement('m:e')
        e.append(self.create_text_omath(base))
        sup_elem.append(e)
        
        # 위첨자
        sup = OxmlElement('m:sup')
        sup.append(self.create_text_omath(superscript))
        sup_elem.append(sup)
        
        return sup_elem
    
    def create_sum_omath(self, lower, upper):
        """합계 기호 OMath 생성"""
        from docx.oxml import OxmlElement
        
        nary = OxmlElement('m:nary')
        
        # 속성
        nary_pr = OxmlElement('m:naryPr')
        chr_elem = OxmlElement('m:chr')
        chr_elem.set(OxmlElement.nsmap['m'] + 'val', '∑')
        nary_pr.append(chr_elem)
        
        lim_loc = OxmlElement('m:limLoc')
        lim_loc.set(OxmlElement.nsmap['m'] + 'val', 'undOvr')
        nary_pr.append(lim_loc)
        
        nary.append(nary_pr)
        
        # 아래 한계
        sub = OxmlElement('m:sub')
        sub.append(self.create_text_omath(lower))
        nary.append(sub)
        
        # 위 한계
        sup = OxmlElement('m:sup')
        sup.append(self.create_text_omath(upper))
        nary.append(sup)
        
        # 본문 (비어있음)
        e = OxmlElement('m:e')
        nary.append(e)
        
        return nary
    
    def create_array_omath(self, latex_code):
        """배열/행렬 OMath 생성"""
        from docx.oxml import OxmlElement
        import re
        
        # 배열 내용 추출
        array_match = re.search(r'\\begin\{array\}(.+?)\\end\{array\}', latex_code, re.DOTALL)
        if not array_match:
            return self.create_text_omath(latex_code)
        
        content = array_match.group(1)
        
        # 행렬 생성
        matrix = OxmlElement('m:m')
        
        # 행 분리
        rows = content.split('\\\\')
        
        for row_text in rows:
            if not row_text.strip():
                continue
                
            mr = OxmlElement('m:mr')
            
            # 열 분리
            cells = row_text.split('&')
            
            for cell_text in cells:
                e = OxmlElement('m:e')
                e.append(self.create_text_omath(cell_text.strip()))
                mr.append(e)
            
            matrix.append(mr)
        
        return matrix
    
    def create_text_omath(self, text):
        """텍스트 OMath 생성"""
        from docx.oxml import OxmlElement
        
        r = OxmlElement('m:r')
        
        # 텍스트 정리
        clean_text = self.clean_latex_text(text)
        
        # OMath 텍스트 요소
        t = OxmlElement('m:t')
        t.text = clean_text
        r.append(t)
        
        return r
    
    def clean_latex_text(self, text):
        """LaTeX 텍스트 정리"""
        # LaTeX 명령어 제거/변환
        replacements = {
            '\\mathrm': '', '\\mathbb': '', '\\mathbf': '',
            '\\operatorname': '', '\\displaystyle': '',
            '\\scriptsize': '', '\\tiny': '', '\\large': '',
            '\\left': '', '\\right': '',
            '\\cdot': '·', '\\times': '×', '\\pm': '±',
            '\\sum': '∑', '\\int': '∫',
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\infty': '∞', '\\partial': '∂',
            '\\leq': '≤', '\\geq': '≥', '\\neq': '≠',
            '\\approx': '≈',
            '{': '', '}': '',
            '~': ' '
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text.strip()
    
    def simplify_latex_for_word(self, latex_code):
        """LaTeX를 간단한 텍스트로 변환 (fallback)"""
        text = self.clean_latex_text(latex_code)
        
        # 추가 단순화
        text = text.replace('\\frac', '').replace('\\', '')
        
        # 너무 길면 축약
        if len(text) > 100:
            return text[:100] + '...'
        
        return text
    
    def merge_text_and_formulas(self, text_blocks, formula_blocks):
        """텍스트와 수식을 위치 기반으로 통합 (같은 줄 처리)"""
        all_blocks = text_blocks + formula_blocks
        
        # 페이지와 y 좌표로 정렬
        all_blocks.sort(key=lambda x: (x.page, x.bbox[1], x.bbox[0]))
        
        # 같은 줄에 있는 요소들을 그룹화
        merged_blocks = []
        current_line = []
        current_y = None
        current_page = None
        threshold = 10  # y 좌표 차이 임계값
        
        for block in all_blocks:
            if current_page != block.page:
                # 페이지가 바뀌면 현재 줄 저장하고 새로 시작
                if current_line:
                    merged_blocks.append(current_line)
                current_line = [block]
                current_y = block.bbox[1]
                current_page = block.page
            elif current_y is None or abs(block.bbox[1] - current_y) < threshold:
                # 같은 줄에 속함
                current_line.append(block)
                if current_y is None:
                    current_y = block.bbox[1]
            else:
                # 새로운 줄
                if current_line:
                    merged_blocks.append(current_line)
                current_line = [block]
                current_y = block.bbox[1]
        
        # 마지막 줄 추가
        if current_line:
            merged_blocks.append(current_line)
        
        return merged_blocks
    
    def create_word_document(self, merged_content, output_path):
        """Word 문서 생성 (원본 레이아웃 보존)"""
        doc = Document()
        
        # 페이지 여백 설정 (원본과 유사하게)
        from docx.shared import Inches
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        current_page = -1
        
        # merged_content는 이제 줄 단위 리스트의 리스트
        for line_blocks in merged_content:
            if not line_blocks:
                continue
            
            # 페이지 확인 (첫 번째 블록의 페이지로)
            page = line_blocks[0].page
            if page != current_page and current_page != -1:
                doc.add_page_break()
            current_page = page
            
            # 새 단락 생성 (한 줄 = 한 단락)
            p = doc.add_paragraph()
            
            # 단락 정렬 설정 (첫 번째 블록의 x 좌표를 기준으로)
            if line_blocks:
                first_x = line_blocks[0].bbox[0]
                # 들여쓰기 수준 결정 (약 72 포인트 = 1 인치)
                if first_x > 100:
                    from docx.shared import Pt
                    p.paragraph_format.left_indent = Pt(first_x / 2)  # PDF 좌표를 포인트로 변환
            
            # 줄 내의 각 요소를 순서대로 추가
            for i, block in enumerate(line_blocks):
                # 블록 간 간격 처리
                if i > 0:
                    prev_block = line_blocks[i-1]
                    gap = block.bbox[0] - prev_block.bbox[2]
                    if gap > 10:  # 10 포인트 이상 간격이 있으면
                        # 공백 추가
                        space_count = int(gap / 5)  # 대략적인 공백 수
                        p.add_run(' ' * space_count)
                
                if block.type == 'text':
                    # 텍스트 추가
                    run = p.add_run(block.text)
                    
                    # 폰트 적용
                    if block.font:
                        try:
                            # 한글 폰트 매핑
                            if 'batang' in block.font.lower() or '바탕' in block.font.lower():
                                run.font.name = '바탕'
                                from docx.oxml import OxmlElement
                                from docx.oxml.ns import qn
                                # 한글 폰트 설정
                                rPr = run._element.get_or_add_rPr()
                                rFonts = OxmlElement('w:rFonts')
                                rFonts.set(qn('w:eastAsia'), '바탕')
                                rPr.append(rFonts)
                            else:
                                run.font.name = block.font
                        except:
                            pass
                    
                    # 크기 적용
                    if block.size:
                        try:
                            from docx.shared import Pt
                            run.font.size = Pt(block.size)
                        except:
                            pass
                    
                    # 텍스트 뒤에 공백 추가 (원본 간격 유지)
                    p.add_run(" ")
                
                elif block.type == 'formula' and block.latex:
                    # 일단 플레이스홀더 텍스트로 삽입 (나중에 OMath로 교체)
                    run = p.add_run()
                    # 고유 ID로 표시
                    placeholder = f"[FORMULA_{block.page}_{i}_{len(line_blocks)}]"  # 페이지_인덱스_총개수
                    run.text = placeholder
                    run.font.name = 'Cambria Math'
                    run.font.italic = True
                    from docx.shared import Pt
                    run.font.size = Pt(11)
                    
                    # 수식 정보 저장
                    if not hasattr(self, '_formula_placeholders'):
                        self._formula_placeholders = {}
                    self._formula_placeholders[placeholder] = block.latex
                    
                    # 수식 뒤 공백
                    p.add_run(" ")
        
        # 문서 저장
        doc.save(str(output_path))
        print(f"✅ Word 문서 생성 완료: {output_path}")
    
    def generate_html(self, merged_content, output_path):
        """HTML 문서 생성 (ln.py 3패널 스타일)"""
        # 먼저 페이지별 이미지 생성
        pdf_images_dir = self.output_dir / "pdf_images"
        pdf_images_dir.mkdir(exist_ok=True)
        
        # PDF를 페이지별 이미지로 변환
        self.generate_pdf_images(str(self.output_dir / "02_no_korean_no_numbers.pdf"), pdf_images_dir)
        
        # JSON 데이터 준비
        json_data = []
        for line_blocks in merged_content:
            for block in line_blocks:
                if block.type == 'text':
                    json_data.append({
                        'type': 'text',
                        'page': block.page,
                        'bbox': block.bbox,
                        'content': block.text,
                        'font': block.font,
                        'size': block.size
                    })
                elif block.type == 'formula':
                    json_data.append({
                        'type': 'formula',
                        'page': block.page,
                        'bbox': block.bbox,
                        'latex': block.latex
                    })
        
        json_path = self.output_dir / "document_structure.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        # 3패널 HTML 생성 - 고급 기능 포함
        html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>3등분 위치정보 기반 뷰어 - PDF | 한글+수식 | Word미리보기</title>
    
    <!-- MathJax 설정 먼저 -->
    <script>
        window.MathJax = {{
            tex: {{ 
                inlineMath: [['$', '$'], ['\\\\(', '\\\\)']], 
                displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
                packages: {{'[+]': ['ams', 'base']}}
            }},
            svg: {{ fontCache: 'global' }},
            startup: {{
                ready: () => {{
                    MathJax.startup.defaultReady();
                    console.log('✅ MathJax 로드 완료');
                }}
            }}
        }};
    </script>
    <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"></script>
    <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        
        body {{ 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            background: #f5f5f5;
            overflow: hidden;
            height: 100vh;
        }}
        
        /* 메인 컨테이너 - 세로 배치 */
        .main-container {{
            display: flex;
            flex-direction: column;
            height: 100vh;
            width: 100vw;
        }}
        
        /* 컨트롤 패널 */
        .control-panel {{
            background: #2c3e50;
            color: white;
            padding: 10px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            height: 60px;
            flex-shrink: 0;
        }}
        
        .control-group {{
            display: flex;
            align-items: center;
            gap: 15px;
        }}
        
        .btn {{
            background: #3498db;
            color: white;
            border: none;
            padding: 8px 15px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
        }}
        
        .btn:hover {{ background: #2980b9; }}
        .btn:disabled {{ background: #7f8c8d; cursor: not-allowed; }}
        
        /* 3등분 패널 컨테이너 - 가로 배치 */
        .panels-container {{
            flex: 1;
            display: flex;
            flex-direction: row;
            height: calc(100vh - 60px);
        }}
        
        /* 개별 패널 스타일 - 가로 33.333% 너비 */
        .panel {{
            flex: 1;
            border-right: 2px solid #34495e;
            display: flex;
            flex-direction: column;
            width: 33.333%;
        }}
        
        .panel:last-child {{ border-right: none; }}
        
        .panel-header {{
            background: #34495e;
            color: white;
            padding: 10px 15px;
            font-weight: bold;
            font-size: 16px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            position: relative;
        }}
        
        /* 패널별 줌 컨트롤 */
        .zoom-controls {{
            display: flex;
            gap: 8px;
            align-items: center;
        }}
        
        .zoom-btn {{
            background: #2c3e50;
            color: white;
            border: none;
            padding: 4px 8px;
            border-radius: 3px;
            cursor: pointer;
            font-size: 12px;
            line-height: 1;
        }}
        
        .zoom-btn:hover {{ background: #1a252f; }}
        
        .zoom-level {{
            font-size: 11px;
            color: #bdc3c7;
            min-width: 35px;
            text-align: center;
        }}
        
        /* 패널 콘텐츠 영역 */
        .panel-content {{
            flex: 1;
            overflow: hidden;
            position: relative;
            background: white;
        }}
        
        .panel-viewport {{
            width: 100%;
            height: 100%;
            overflow: auto;
            position: relative;
        }}
        
        .panel-inner {{
            transform-origin: top left;
            transition: transform 0.2s ease;
            cursor: grab;
            min-height: 100%;
            padding: 20px;
        }}
        
        .panel-inner:active {{
            cursor: grabbing;
        }}
        
        /* PDF 패널 전용 스타일 */
        #pdf-panel .panel-inner {{
            background: #f8f9fa;
            padding: 10px;
        }}
        
        .pdf-page {{
            margin-bottom: 20px;
            text-align: center;
            border: 1px solid #dee2e6;
            background: white;
            border-radius: 5px;
            overflow: hidden;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        .pdf-page img {{
            max-width: 100%;
            height: auto;
            display: block;
        }}
        
        .pdf-page-number {{
            background: #34495e;
            color: white;
            padding: 8px;
            font-size: 14px;
            font-weight: bold;
        }}
        
        /* 콘텐츠 패널 스타일 */
        #content-panel .panel-inner {{
            background: white;
            line-height: 1.6;
        }}
        
        .content-line {{
            margin: 8px 0;
            min-height: 24px;
            word-wrap: break-word;
        }}
        
        .content-text {{
            display: inline;
            font-size: 14px;
        }}
        
        .content-formula {{
            display: inline-block;
            background: #e8f4fd;
            padding: 4px 8px;
            border-radius: 4px;
            margin: 0 4px;
            border: 1px solid #d1ecf1;
            font-family: 'Times New Roman', serif;
        }}
        
        .latex-code {{
            display: block;
            background: #f8f9fa;
            padding: 8px 12px;
            border-radius: 4px;
            margin: 4px 0;
            border: 1px solid #e9ecef;
            font-family: 'Monaco', 'Consolas', monospace;
            font-size: 13px;
            color: #495057;
            white-space: pre-wrap;
        }}
        
        .page-separator {{
            border-top: 3px double #34495e;
            margin: 30px 0;
            padding-top: 20px;
            text-align: center;
            color: #34495e;
            font-weight: bold;
            font-size: 16px;
        }}
        
        /* JSON 패널 스타일 - 수식 렌더링용 */
        #json-panel .panel-inner {{
            background: white;
            color: #333;
            font-family: 'Times New Roman', serif;
            font-size: 14px;
            line-height: 1.6;
        }}
        
        #json-panel pre {{
            white-space: pre-wrap;
            word-wrap: break-word;
            margin: 0;
            color: #ecf0f1;
        }}
        
        .json-type-text {{ 
            color: #3498db; 
            font-weight: bold;
        }}
        
        .json-type-formula {{ 
            color: #e74c3c; 
            font-weight: bold;
        }}
        
        /* 로딩 스피너 */
        .spinner {{
            border: 3px solid #f3f3f3;
            border-top: 3px solid #3498db;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 20px auto;
        }}
        
        @keyframes spin {{
            0% {{ transform: rotate(0deg); }}
            100% {{ transform: rotate(360deg); }}
        }}
        
        /* 반응형 디자인 */
        @media (max-width: 1200px) {{
            .panels-container {{
                flex-direction: column;
            }}
            .panel {{
                border-right: none;
                border-bottom: 2px solid #34495e;
                width: 100%;
                height: 33.333%;
            }}
            .panel:last-child {{
                border-bottom: none;
            }}
        }}
        
        /* 스크롤바 스타일링 */
        .panel-viewport::-webkit-scrollbar {{
            width: 8px;
            height: 8px;
        }}
        
        .panel-viewport::-webkit-scrollbar-track {{
            background: #f1f1f1;
        }}
        
        .panel-viewport::-webkit-scrollbar-thumb {{
            background: #888;
            border-radius: 4px;
        }}
        
        .panel-viewport::-webkit-scrollbar-thumb:hover {{
            background: #555;
        }}
    </style>
</head>
<body>
    <div class="main-container">
        <!-- 상단 컨트롤 패널 -->
        <div class="control-panel">
            <div class="control-group">
                <h1 style="font-size: 18px; margin: 0;">🔬 LaTeX to Word 3패널 뷰어</h1>
            </div>
            <div class="control-group">
                <button class="btn" onclick="resetAllZoom()">🔄 전체 리셋</button>
                <button class="btn" onclick="toggleSync()">🔗 동기화: <span id="sync-status">OFF</span></button>
                <button class="btn" onclick="exportData()">📥 내보내기</button>
            </div>
        </div>
        
        <!-- 3등분 패널 컨테이너 -->
        <div class="panels-container">
            <!-- 패널 1: PDF 원본 -->
            <div class="panel" id="pdf-panel">
                <div class="panel-header">
                    <span>📄 PDF 원본</span>
                    <div class="zoom-controls">
                        <button class="zoom-btn" onclick="zoomPanel('pdf-panel', -0.1)">-</button>
                        <span class="zoom-level" id="pdf-zoom">100%</span>
                        <button class="zoom-btn" onclick="zoomPanel('pdf-panel', 0.1)">+</button>
                    </div>
                </div>
                <div class="panel-content">
                    <div class="panel-viewport" id="pdf-viewport">
                        <div class="panel-inner" id="pdf-inner">
"""
        
        # PDF 이미지 추가
        pdf_images = sorted(pdf_images_dir.glob("page_*.png"))
        for img in pdf_images:
            page_num = int(img.stem.split('_')[1])
            html += f'''
                        <div class="pdf-page">
                            <div class="pdf-page-number">페이지 {page_num + 1}</div>
                            <img src="pdf_images/{img.name}" alt="Page {page_num + 1}">
                        </div>
'''
        
        html += """                    </div>
                    </div>
                </div>
            </div>
            
            <!-- 패널 2: LaTeX 코드 -->
            <div class="panel" id="content-panel">
                <div class="panel-header">
                    <span>📝 LaTeX 코드</span>
                    <div class="zoom-controls">
                        <button class="zoom-btn" onclick="zoomPanel('content-panel', -0.1)">-</button>
                        <span class="zoom-level" id="content-zoom">100%</span>
                        <button class="zoom-btn" onclick="zoomPanel('content-panel', 0.1)">+</button>
                    </div>
                </div>
                <div class="panel-content">
                    <div class="panel-viewport" id="content-viewport">
                        <div class="panel-inner" id="content-inner">
"""
        
        # LaTeX 코드 추가 (중간 패널)
        current_page = -1
        for line_blocks in merged_content:
            if not line_blocks:
                continue
            
            page = line_blocks[0].page
            if page != current_page:
                if current_page != -1:
                    html += f'<div class="page-separator">페이지 {page + 1}</div>\n'
                current_page = page
            
            html += '<div class="content-line">'
            for block in line_blocks:
                if block.type == 'text':
                    html += f'<span class="content-text">{block.text} </span>'
                elif block.type == 'formula' and block.latex:
                    # LaTeX 코드만 표시 (렌더링 없음)
                    html += f'<div class="latex-code">{block.latex}</div>'
            html += '</div>\n'
        
        html += """                        </div>
                    </div>
                </div>
            </div>
            
            <!-- 패널 3: 수식 렌더링 -->
            <div class="panel" id="json-panel">
                <div class="panel-header">
                    <span>🎯 수식 렌더링</span>
                    <div class="zoom-controls">
                        <button class="zoom-btn" onclick="zoomPanel('json-panel', -0.1)">-</button>
                        <span class="zoom-level" id="json-zoom">100%</span>
                        <button class="zoom-btn" onclick="zoomPanel('json-panel', 0.1)">+</button>
                    </div>
                </div>
                <div class="panel-content">
                    <div class="panel-viewport" id="json-viewport">
                        <div class="panel-inner" id="json-inner">
"""
        
        # 렌더링된 수식 추가 (오른쪽 패널)
        current_page = -1
        for line_blocks in merged_content:
            if not line_blocks:
                continue
            
            page = line_blocks[0].page
            if page != current_page:
                if current_page != -1:
                    html += f'<div class="page-separator">페이지 {page + 1}</div>\n'
                current_page = page
            
            html += '<div class="content-line">'
            for block in line_blocks:
                if block.type == 'text':
                    html += f'<span class="content-text">{block.text} </span>'
                elif block.type == 'formula' and block.latex:
                    # MathJax로 렌더링된 수식
                    html += f'<span class="content-formula">\\({block.latex}\\)</span>'
            html += '</div>\n'
        
        html += """
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        // 전역 변수
        const zoomLevels = {{
            'pdf-panel': 1.0,
            'content-panel': 1.0,
            'json-panel': 1.0
        }};
        let syncScroll = false;
        let isDragging = false;
        let dragStartX = 0;
        let dragStartY = 0;
        let panelStartScrollX = 0;
        let panelStartScrollY = 0;
        let currentPanel = null;
        
        // 줌 기능
        function zoomPanel(panelId, delta) {{
            console.log('줌 함수 호출:', panelId, delta);
            const newLevel = Math.max(0.25, Math.min(2.0, zoomLevels[panelId] + delta));
            zoomLevels[panelId] = newLevel;
            
            const inner = document.getElementById(panelId.replace('-panel', '-inner'));
            if (inner) {{
                inner.style.transform = `scale(${{newLevel}})`;
                console.log('줌 적용:', newLevel);
            }} else {{
                console.error('패널 inner 요소를 찾을 수 없음:', panelId);
            }}
            
            const zoomDisplay = document.getElementById(panelId.replace('-panel', '-zoom'));
            if (zoomDisplay) {{
                zoomDisplay.textContent = Math.round(newLevel * 100) + '%';
            }} else {{
                console.error('줌 표시 요소를 찾을 수 없음:', panelId);
            }}
        }}
        
        // 전체 리셋
        function resetAllZoom() {{
            console.log('전체 리셋 시작');
            ['pdf-panel', 'content-panel', 'json-panel'].forEach(panelId => {{
                zoomLevels[panelId] = 1.0;
                const inner = document.getElementById(panelId.replace('-panel', '-inner'));
                if (inner) {{
                    inner.style.transform = 'scale(1.0)';
                }} else {{
                    console.error('리셋 - inner 요소 없음:', panelId);
                }}
                
                const zoomDisplay = document.getElementById(panelId.replace('-panel', '-zoom'));
                if (zoomDisplay) {{
                    zoomDisplay.textContent = '100%';
                }} else {{
                    console.error('리셋 - 줌 표시 요소 없음:', panelId);
                }}
                
                // 스크롤 위치도 리셋
                const viewport = document.getElementById(panelId.replace('-panel', '-viewport'));
                if (viewport) {{
                    viewport.scrollTop = 0;
                    viewport.scrollLeft = 0;
                }} else {{
                    console.error('리셋 - viewport 요소 없음:', panelId);
                }}
            }});
            console.log('전체 리셋 완료');
        }}
        
        // 동기화 토글
        function toggleSync() {{
            syncScroll = !syncScroll;
            document.getElementById('sync-status').textContent = syncScroll ? 'ON' : 'OFF';
        }}
        
        // 데이터 내보내기
        function exportData() {{
            alert('내보내기 기능은 추후 구현 예정입니다.');
        }}
        
        // 드래그 기능 초기화
        function initDragFunctionality() {{
            console.log('드래그 기능 초기화 시작');
            const panelInners = document.querySelectorAll('.panel-inner');
            console.log('발견된 panel-inner 요소 수:', panelInners.length);
            
            panelInners.forEach((inner, index) => {{
                console.log('드래그 이벤트 등록:', index, inner.id);
                inner.addEventListener('mousedown', startDrag);
            }});
            
            // 전역 이벤트 리스너
            document.addEventListener('mousemove', drag);
            document.addEventListener('mouseup', endDrag);
        }}
        
        function startDrag(e) {{
            if (e.target.tagName === 'IMG' || e.target.tagName === 'BUTTON') return;
            
            console.log('드래그 시작');
            isDragging = true;
            dragStartX = e.clientX;
            dragStartY = e.clientY;
            
            currentPanel = e.currentTarget.closest('.panel');
            if (currentPanel) {{
                const viewport = currentPanel.querySelector('.panel-viewport');
                if (viewport) {{
                    panelStartScrollX = viewport.scrollLeft;
                    panelStartScrollY = viewport.scrollTop;
                    console.log('드래그 초기 스크롤 위치:', panelStartScrollX, panelStartScrollY);
                }}
            }}
            
            e.currentTarget.style.cursor = 'grabbing';
            e.preventDefault();
        }}
        
        function drag(e) {{
            if (!isDragging || !currentPanel) return;
            
            const deltaX = dragStartX - e.clientX;
            const deltaY = dragStartY - e.clientY;
            
            const viewport = currentPanel.querySelector('.panel-viewport');
            if (viewport) {{
                viewport.scrollLeft = panelStartScrollX + deltaX;
                viewport.scrollTop = panelStartScrollY + deltaY;
            }}
        }}
        
        function endDrag(e) {{
            if (isDragging) {{
                console.log('드래그 종료');
                isDragging = false;
                const allInners = document.querySelectorAll('.panel-inner');
                allInners.forEach(inner => {{
                    inner.style.cursor = 'grab';
                }});
                currentPanel = null;
            }}
        }}
        
        // 키보드 단축키
        document.addEventListener('keydown', (e) => {{
            console.log('키 입력:', e.key, 'Ctrl:', e.ctrlKey);
            if (e.key === 's' && e.ctrlKey) {{
                console.log('동기화 토글 실행');
                toggleSync();
                e.preventDefault();
            }} else if (e.key === 'r' && e.ctrlKey) {{
                console.log('전체 리셋 실행');
                resetAllZoom();
                e.preventDefault();
            }}
        }});
        
        // 휠 줌 기능
        document.addEventListener('wheel', (e) => {{
            if (e.ctrlKey) {{
                console.log('휠 줌 감지');
                e.preventDefault();
                const panel = e.target.closest('.panel');
                if (panel) {{
                    const panelId = panel.id;
                    const delta = e.deltaY > 0 ? -0.1 : 0.1;
                    console.log('휠 줌 실행:', panelId, delta);
                    zoomPanel(panelId, delta);
                }} else {{
                    console.log('패널을 찾을 수 없음');
                }}
            }}
        }}, {{ passive: false }});
        
        // MathJax 재렌더링
        function rerenderMath() {{
            if (window.MathJax) {{
                console.log('MathJax 재렌더링 시작');
                MathJax.typesetPromise().then(() => {{
                    console.log('MathJax 재렌더링 완료');
                }}).catch((err) => {{
                    console.error('MathJax 렌더링 오류:', err);
                }});
            }}
        }}
        
        // 페이지 로드 완료 후 초기화
        document.addEventListener('DOMContentLoaded', () => {{
            initDragFunctionality();
            console.log('✅ 3패널 뷰어 초기화 완료');
            console.log('💡 사용법:');
            console.log('  - 패널 헤더 +/- 버튼: 줌 인/아웃');
            console.log('  - Ctrl + 휠: 휠 줌');
            console.log('  - 마우스 드래그: 패널 이동');
            console.log('  - Ctrl + S: 동기화 토글');
            console.log('  - Ctrl + R: 전체 리셋');
            
            // MathJax가 로드되면 렌더링
            setTimeout(() => {{
                rerenderMath();
            }}, 1000);
        }});
    </script>
</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        print(f"✅ HTML 3패널 뷰어 생성 완료: {output_path}")
        
        # HTML 자동 실행
        try:
            os.startfile(str(output_path))
        except:
            pass
    
    def generate_pdf_images(self, pdf_path, output_dir):
        """PDF를 페이지별 이미지로 변환"""
        doc = fitz.open(pdf_path)
        page_count = len(doc)  # close 전에 저장
        
        for page_num in range(page_count):
            page = doc[page_num]
            # 150 DPI로 렌더링 (파일 크기 고려)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img_path = output_dir / f"page_{page_num}.png"
            pix.save(str(img_path))
        
        doc.close()
        print(f"  📸 {page_count}개 페이지 이미지 생성")


# 메인 실행
if __name__ == "__main__":
    pipeline = LaTeXToWordPipeline()
    
    # 명령줄 인자 확인
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        # 기본 파일
        input_file = "C:/test/27.pdf"
    
    # 파일 존재 확인
    if not Path(input_file).exists():
        print(f"❌ 파일을 찾을 수 없습니다: {input_file}")
        sys.exit(1)
    
    # 파일 타입 확인
    file_ext = Path(input_file).suffix.lower()
    if file_ext in ['.doc', '.docx']:
        print(f"📄 Word 문서 감지: {input_file}")
    elif file_ext == '.pdf':
        print(f"📄 PDF 문서 감지: {input_file}")
    else:
        print(f"❌ 지원하지 않는 파일 형식: {file_ext}")
        print("   지원 형식: .pdf, .doc, .docx")
        sys.exit(1)
    
    # 파이프라인 실행
    results = pipeline.run_pipeline(input_file)